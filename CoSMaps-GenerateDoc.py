# -*- coding: utf-8 -*-
"""
Created on Fri Jan  1 16:43:34 2021

@author: StormyWeather
"""

# import datetime as dt
from openpyxl import load_workbook
import docx
from docx import Document

class useful_link:
    def __init__(self, title, date, resource, post, useful, tag1, tag2, tag3):
        self.title = title
        self.date = date
        self.resource = resource
        self.post = post
        self.useful = useful
        self.tag1 = tag1
        self.tag2 = tag2
        self.tag3 = tag3


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # StormyWeather edit: Make font Arial
    f = docx.oxml.shared.OxmlElement('w:rFonts')
    f.set(docx.oxml.shared.qn('w:hAnsi'),'Arial')
    f.set(docx.oxml.shared.qn('w:ascii'),'Arial')
    rPr.append(f)

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    # if not underline:
    #   u = docx.oxml.shared.OxmlElement('w:u')
    #   u.set(docx.oxml.shared.qn('w:val'), 'none')
    #   rPr.append(u)

    # StormyWeather edit: Underline all links
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

# Open the spreadsheet
file = 'CoSMapsPSAW.xlsx'
wb = load_workbook(filename = file)
ws = wb['CoSMapsPSAW']

# Create empty list of link objects
links = []
count = 0

# read the source workbook and parse into class objects
# tag2 is defaulted to 'A' so that things not sorted by subsection go before
# everything else. Kind of a hack but it works.
for row in ws.iter_rows(min_row=2):
    if (row[4].value == 'y'):
        link = useful_link(row[0].value,
                           row[1].value,
                           row[2].value,
                           row[3].value,
                           row[4].value,
                           row[5].value,
                           'A',
                           None)
        if (row[6].value):
            link.tag2 = row[6].value
        if (row[7].value):
            link.tag3 = row[7].value
        links.append(link)
        count += 1
    else:
        continue

doc = Document()

# Write the intro paragraphs
p = doc.add_paragraph()
run = p.add_run("# r/CoS Maps")
font = run.font
font.name = 'Arial'
last_updated = ws['J1'].value
last_updated = last_updated[:10]

p = doc.add_paragraph()
run = p.add_run("Last updated: " + last_updated)
font = run.font
font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("This page attempts to be an organized, master resource list for all the maps "
                      + "submitted to r/CurseofStrahd. The focus of this list is on resources that could "
                      + "be used by DMs. I have tried to be generous when it comes to art style and quality "
                      + "since every DM has their own tastes and preferences.")
font = run.font
font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("This list is organized first by larger areas (e.g. Vallaki) and then by sub-areas as needed "
                      +"(e.g. Wachterhaus). This list also attempts to include maps for popular expansions such as "
                      +"CoS: Reloaded. If you are a map creator and feel I have missed one of your submissions, please "
                      +"contact the mods and they will pass the message along to me.")
font = run.font
font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run("In some cases, you will have to look through the comments on a post to find a link to the resource.")
font = run.font
font.name = 'Arial'


# section and subsection are used to sort links into the correct paragraphs
section = None
subsection = None

# sort the links by section, subsection, and then title
links = sorted(links, key=lambda x: (x.tag1, x.tag2, x.title))

# for every link in link
for link in links:
    # skip_newline is used to catch a special case where a section that has
    # nothing tagged 'A' (aka, no subsection) will get some weird extra
    # new lines. It just looks weird. skip_newline is another kind of hack
    # but it works so whatever.
    skip_newline = False
    
    # if the current section is different from this link, start a new paragraph
    if section != link.tag1:
        
        # check for the weird case I described above
        if section != link.tag1 and link.tag2 != 'A':
            skip_newline = True
        
        # set section to the correct current value
        section = link.tag1
        
        # add a new paragraph for the section
        p = doc.add_paragraph()
        run = p.add_run("## " + str(link.tag1).rstrip())
        font = run.font
        font.name = 'Arial'
        
        # if that weird case I described above happens, don't print a new line
        if not skip_newline:
            run = p.add_run("\n")
    
    # if the current subsection is different from this link and it's not an
    # unsorted link ('A' is used to put links with no subsection first)
    # the start a new paragraph
    if link.tag2 != subsection and link.tag2 != 'A':
        
        # set subsection to the correct current value
        subsection = link.tag2
        
        # add a new paragraph for the subsection
        p = doc.add_paragraph()
        run = p.add_run("### " + str(link.tag2))
        font = run.font
        font.name = 'Arial'
        run = p.add_run("\n")
    
    # add the hyperlink
    hyperlink = add_hyperlink(p, link.post, link.title, '0000EE', True)
    
    # add the date of the link to the line
    # run = p.add_run(" - " + link.date[:10])
    # font = run.font
    # font.name = 'Arial'
    
    # If there is anything in 'tag3' add it to the line, aka Editor's Choice
    # or Fleshing Out Curse of Strahd
    if link.tag3 != None:
        run = p.add_run(" - " + link.tag3)
        run.bold = True
        font = run.font
        font.name = 'Arial'
    run = p.add_run("\n")
    
    # print(link.tag1 + ':' + link.tag2)
print("Lines added: {}".format(count))


doc.save("CoS - Wiki - Maps.docx")
